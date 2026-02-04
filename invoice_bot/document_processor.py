#!/usr/bin/env python3
"""
Document processing module for extracting invoice data from uploaded files.
Supports PDF, DOCX, and image files (JPEG, PNG).

SELF-HEALING FEATURES:
- All blocking OCR calls run in thread pool (asyncio.to_thread)
- Timeout handling for all operations
- Fallback to rule-based extraction on failure
- Detailed logging for debugging
"""

import asyncio
import os
import re
import logging
from typing import Dict, Any, Optional, List
from datetime import datetime
from functools import partial

from invoice_bot.invoice_data import InvoiceData, WorkItem
from invoice_bot.ai_integration import ai_assessor, QUALITY_THRESHOLD

logger = logging.getLogger(__name__)

# Timeouts for different operations
OCR_TIMEOUT = 30  # seconds
PDF_TIMEOUT = 45  # seconds
DOCX_TIMEOUT = 15  # seconds

# Quality thresholds
MIN_QUALITY_THRESHOLD = QUALITY_THRESHOLD  # From ai_integration


class DocumentProcessor:
    """Process uploaded documents and extract invoice data."""
    
    def __init__(self):
        self.patterns = self._compile_patterns()
    
    def _compile_patterns(self) -> Dict[str, re.Pattern]:
        """Compile regex patterns for data extraction."""
        return {
            # Email pattern
            "email": re.compile(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'),
            
            # UK Postcode pattern
            "postcode": re.compile(r'[A-Z]{1,2}[0-9][A-Z0-9]?\s*[0-9][A-Z]{2}'),
            
            # Phone pattern
            "phone": re.compile(r'(?:\+44|0)\s?[0-9]{4}\s?[0-9]{6}'),
            
            # UTR pattern (10 digits)
            "utr": re.compile(r'\b\d{10}\b'),
            
            # NI Number pattern
            "ni_number": re.compile(r'[A-CEGHJ-PR-TW-Z]{2}\s?\d{2}\s?\d{2}\s?\d{2}\s?[A-D]'),
            
            # Sort code pattern
            "sort_code": re.compile(r'\b\d{2}[-\s]?\d{2}[-\s]?\d{2}\b'),
            
            # Bank account pattern (8 digits)
            "bank_account": re.compile(r'\b\d{8}\b'),
            
            # Invoice number patterns
            "invoice_number": re.compile(r'(?:invoice\s*(?:#|number|no)?[:\s]*)?([A-Z0-9-]{3,20})', re.IGNORECASE),
            
            # Date patterns (DD/MM/YYYY or variations)
            "date": re.compile(r'\b(\d{1,2})[/-](\d{1,2})[/-](\d{2,4})\b'),
            
            # Currency amounts
            "amount": re.compile(r'£?\s*(\d{1,3}(?:,\d{3})*\.?\d{0,2})'),
            
            # VAT amount
            "vat": re.compile(r'(?:vat|tax)[:\s]*£?\s*(\d+\.?\d{0,2})', re.IGNORECASE),
            
            # CIS deduction
            "cis": re.compile(r'(?:cis|deduction)[:\s]*£?\s*(\d+\.?\d{0,2})', re.IGNORECASE),
            
            # Subtotal/Total patterns
            "subtotal": re.compile(r'(?:subtotal|sub-total|net)[:\s]*£?\s*(\d+\.?\d{0,2})', re.IGNORECASE),
            "total": re.compile(r'(?:total|amount\s*due|grand\s*total)[:\s]*£?\s*(\d+\.?\d{0,2})', re.IGNORECASE),
        }
    
    async def process_document(self, file_path: str, mime_type: str) -> InvoiceData:
        """
        Process a document and extract invoice data.
        
        Args:
            file_path: Path to the downloaded file
            mime_type: MIME type of the file
            
        Returns:
            InvoiceData object with extracted information
        """
        logger.info(f"Processing document: {file_path}, MIME: {mime_type}")
        
        # Extract text from document
        text = await self._extract_text(file_path, mime_type)
        
        if not text:
            logger.warning("No text extracted from document")
            return InvoiceData()
        
        logger.info(f"Extracted text length: {len(text)}")

        # Parse invoice data from text
        invoice_data = self._parse_invoice_data(text)

        return invoice_data

    async def process_document_with_healing(self, file_path: str, mime_type: str) -> InvoiceData:
        """
        Self-healing document processing with AI quality assessment.

        This method:
        1. Tries standard OCR extraction
        2. Assesses quality with AI
        3. If quality is poor, tries vision API fallback
        4. Only returns data if quality threshold is met
        5. Marks extraction as failed if all strategies fail
        """
        logger.info(f"[SELF-HEALING] Starting extraction: {file_path}")

        # Strategy 1: Standard OCR extraction
        logger.info("[SELF-HEALING] Strategy 1: Standard OCR extraction")
        invoice_data = await self.process_document(file_path, mime_type)
        data_dict = invoice_data.to_dict()

        # Assess quality
        quality = await ai_assessor.assess_quality(data_dict)
        logger.info(f"[SELF-HEALING] OCR quality score: {quality:.2f}")

        if quality >= MIN_QUALITY_THRESHOLD:
            logger.info(f"[SELF-HEALING] Quality acceptable ({quality:.2f} >= {MIN_QUALITY_THRESHOLD})")
            invoice_data.extraction_confidence = quality
            return invoice_data

        # Strategy 2: Enhanced OCR with preprocessing
        if mime_type.startswith("image/"):
            logger.info("[SELF-HEALING] Strategy 2: Enhanced OCR with preprocessing")
            enhanced_text = await self._extract_with_preprocessing(file_path)
            if enhanced_text:
                invoice_data = self._parse_invoice_data(enhanced_text)
                data_dict = invoice_data.to_dict()
                quality = await ai_assessor.assess_quality(data_dict)
                logger.info(f"[SELF-HEALING] Enhanced OCR quality: {quality:.2f}")

                if quality >= MIN_QUALITY_THRESHOLD:
                    invoice_data.extraction_confidence = quality
                    return invoice_data

        # Strategy 3: AI Vision extraction (GPT-4 Vision)
        if mime_type.startswith("image/"):
            logger.info("[SELF-HEALING] Strategy 3: AI Vision extraction")
            vision_data = await ai_assessor.extract_with_vision(file_path)

            if vision_data:
                # Convert vision data to InvoiceData
                invoice_data = self._vision_data_to_invoice(vision_data)
                data_dict = invoice_data.to_dict()
                quality = await ai_assessor.assess_quality(data_dict)
                logger.info(f"[SELF-HEALING] Vision extraction quality: {quality:.2f}")

                if quality >= MIN_QUALITY_THRESHOLD:
                    invoice_data.extraction_confidence = quality
                    invoice_data.source = "ai_vision"
                    return invoice_data

        # All strategies failed
        logger.warning(f"[SELF-HEALING] All extraction strategies failed (best quality: {quality:.2f})")
        invoice_data.extraction_failed = True
        invoice_data.extraction_confidence = quality
        return invoice_data

    async def _extract_with_preprocessing(self, file_path: str) -> str:
        """
        Extract text with image preprocessing for better OCR.

        Applies:
        - Grayscale conversion
        - Contrast enhancement
        - Noise reduction
        - Thresholding
        """
        try:
            import pytesseract
            from PIL import Image, ImageEnhance, ImageFilter

            def _sync_enhanced_ocr():
                # Open and preprocess image
                image = Image.open(file_path)

                # Convert to grayscale
                if image.mode != 'L':
                    image = image.convert('L')

                # Enhance contrast
                enhancer = ImageEnhance.Contrast(image)
                image = enhancer.enhance(2.0)

                # Apply sharpening
                image = image.filter(ImageFilter.SHARPEN)

                # Apply threshold for cleaner text
                threshold = 150
                image = image.point(lambda p: 255 if p > threshold else 0)

                # OCR with optimized settings
                custom_config = r'--oem 3 --psm 6'
                return pytesseract.image_to_string(image, config=custom_config)

            logger.info("Running enhanced OCR with preprocessing...")
            text = await asyncio.wait_for(
                asyncio.to_thread(_sync_enhanced_ocr),
                timeout=OCR_TIMEOUT
            )
            logger.info(f"Enhanced OCR extracted {len(text)} characters")
            return text

        except Exception as e:
            logger.error(f"Enhanced OCR failed: {e}")
            return ""

    def _vision_data_to_invoice(self, vision_data: dict) -> InvoiceData:
        """Convert AI Vision extraction result to InvoiceData object."""
        invoice = InvoiceData()
        invoice.source = "ai_vision"

        # Map vision data to invoice fields
        invoice.contractor_name = vision_data.get('contractor_name') or ""
        invoice.contractor_email = vision_data.get('contractor_email') or ""
        invoice.contractor_address = vision_data.get('contractor_address') or ""
        invoice.contractor_utr = vision_data.get('contractor_utr') or ""
        invoice.contractor_ni = vision_data.get('contractor_ni') or ""
        invoice.bank_account = vision_data.get('bank_account') or ""
        invoice.sort_code = vision_data.get('sort_code') or ""
        invoice.invoice_number = vision_data.get('invoice_number') or ""
        invoice.invoice_date = vision_data.get('invoice_date') or ""
        invoice.work_start_date = vision_data.get('work_start_date') or ""
        invoice.work_end_date = vision_data.get('work_end_date') or ""

        # Financial data
        invoice.subtotal = float(vision_data.get('subtotal') or 0)
        invoice.vat_amount = float(vision_data.get('vat_amount') or 0)
        invoice.cis_amount = float(vision_data.get('cis_amount') or 0)
        invoice.total = float(vision_data.get('total') or 0)

        if invoice.total == 0:
            invoice.calculate_total()

        return invoice
    
    async def _extract_text(self, file_path: str, mime_type: str) -> str:
        """Extract text from various document types."""
        try:
            if mime_type == "application/pdf":
                return await self._extract_from_pdf(file_path)
            elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                return await self._extract_from_docx(file_path)
            elif mime_type.startswith("image/"):
                return await self._extract_from_image(file_path)
            else:
                logger.warning(f"Unsupported MIME type: {mime_type}")
                return ""
        except Exception as e:
            logger.error(f"Text extraction error: {e}")
            return ""
    
    async def _extract_from_pdf(self, file_path: str) -> str:
        """
        Extract text from PDF file.

        SELF-HEALING: Multiple fallback methods with timeout protection.
        """
        # Try PyPDF2 first (fast, text-based PDFs)
        try:
            from PyPDF2 import PdfReader

            def _sync_pypdf():
                reader = PdfReader(file_path)
                text = ""
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                return text

            logger.info("Trying PyPDF2 extraction...")
            text = await asyncio.wait_for(
                asyncio.to_thread(_sync_pypdf),
                timeout=PDF_TIMEOUT
            )
            if text.strip():
                logger.info(f"PyPDF2 extracted {len(text)} characters")
                return text
        except ImportError:
            logger.warning("PyPDF2 not available")
        except asyncio.TimeoutError:
            logger.warning("PyPDF2 timed out")
        except Exception as e:
            logger.warning(f"PyPDF2 failed: {e}")

        # Fallback to pdfplumber
        try:
            import pdfplumber

            def _sync_pdfplumber():
                text = ""
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                return text

            logger.info("Trying pdfplumber extraction...")
            text = await asyncio.wait_for(
                asyncio.to_thread(_sync_pdfplumber),
                timeout=PDF_TIMEOUT
            )
            if text.strip():
                logger.info(f"pdfplumber extracted {len(text)} characters")
                return text
        except ImportError:
            logger.warning("pdfplumber not available")
        except asyncio.TimeoutError:
            logger.warning("pdfplumber timed out")
        except Exception as e:
            logger.warning(f"pdfplumber failed: {e}")

        # Last resort: OCR with pdf2image (slow, for scanned PDFs)
        try:
            from pdf2image import convert_from_path
            import pytesseract

            def _sync_pdf_ocr():
                images = convert_from_path(file_path)
                text = ""
                for i, image in enumerate(images):
                    logger.info(f"OCR processing page {i+1}/{len(images)}")
                    text += pytesseract.image_to_string(image) + "\n"
                return text

            logger.info("Trying PDF OCR extraction (this may take a while)...")
            text = await asyncio.wait_for(
                asyncio.to_thread(_sync_pdf_ocr),
                timeout=PDF_TIMEOUT * 2  # Double timeout for OCR
            )
            logger.info(f"PDF OCR extracted {len(text)} characters")
            return text
        except ImportError:
            logger.error("No PDF extraction library available")
        except asyncio.TimeoutError:
            logger.error(f"PDF OCR timed out after {PDF_TIMEOUT * 2}s")
        except Exception as e:
            logger.error(f"PDF OCR failed: {e}")

        return ""
    
    async def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        try:
            from docx import Document
            doc = Document(file_path)
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text
        except ImportError:
            logger.error("python-docx not available")
            return ""
    
    async def _extract_from_image(self, file_path: str) -> str:
        """
        Extract text from image using OCR.

        SELF-HEALING: Runs in thread pool with timeout to prevent blocking.
        """
        try:
            import pytesseract
            from PIL import Image

            def _sync_ocr():
                """Synchronous OCR operation to run in thread pool."""
                image = Image.open(file_path)
                return pytesseract.image_to_string(image)

            # Run blocking OCR in thread pool with timeout
            logger.info(f"Starting OCR for {file_path} (timeout: {OCR_TIMEOUT}s)")
            text = await asyncio.wait_for(
                asyncio.to_thread(_sync_ocr),
                timeout=OCR_TIMEOUT
            )
            logger.info(f"OCR completed, extracted {len(text)} characters")
            return text

        except asyncio.TimeoutError:
            logger.error(f"OCR timed out after {OCR_TIMEOUT}s for {file_path}")
            # Return empty string - orchestrator will handle fallback
            return ""
        except ImportError:
            logger.error("pytesseract or PIL not available")
            return ""
        except Exception as e:
            logger.error(f"OCR error: {e}")
            return ""
    
    def _parse_invoice_data(self, text: str) -> InvoiceData:
        """Parse invoice data from extracted text."""
        invoice = InvoiceData()
        invoice.source = "upload"
        
        lines = text.split('\n')
        
        # Extract email
        email_match = self.patterns["email"].search(text)
        if email_match:
            invoice.contractor_email = email_match.group(0)
        
        # Extract UTR
        utr_match = self.patterns["utr"].search(text)
        if utr_match:
            invoice.contractor_utr = utr_match.group(0)
        
        # Extract NI number
        ni_match = self.patterns["ni_number"].search(text)
        if ni_match:
            invoice.contractor_ni = ni_match.group(0)
        
        # Extract sort code
        sort_match = self.patterns["sort_code"].search(text)
        if sort_match:
            invoice.sort_code = sort_match.group(0)
        
        # Extract bank account
        bank_match = self.patterns["bank_account"].search(text)
        if bank_match:
            invoice.bank_account = bank_match.group(0)
        
        # Extract invoice number
        inv_match = self.patterns["invoice_number"].search(text)
        if inv_match:
            invoice.invoice_number = inv_match.group(1)
        
        # Extract dates
        dates = self.patterns["date"].findall(text)
        if len(dates) >= 1:
            invoice.invoice_date = f"{dates[0][0]}/{dates[0][1]}/{dates[0][2]}"
        if len(dates) >= 2:
            invoice.work_start_date = f"{dates[1][0]}/{dates[1][1]}/{dates[1][2]}"
        if len(dates) >= 3:
            invoice.work_end_date = f"{dates[2][0]}/{dates[2][1]}/{dates[2][2]}"
        
        # Extract financial amounts
        subtotal_match = self.patterns["subtotal"].search(text)
        if subtotal_match:
            invoice.subtotal = float(subtotal_match.group(1).replace(',', ''))
        
        vat_match = self.patterns["vat"].search(text)
        if vat_match:
            invoice.vat_amount = float(vat_match.group(1).replace(',', ''))
        
        cis_match = self.patterns["cis"].search(text)
        if cis_match:
            invoice.cis_amount = float(cis_match.group(1).replace(',', ''))
        
        total_match = self.patterns["total"].search(text)
        if total_match:
            invoice.total = float(total_match.group(1).replace(',', ''))
        
        # Try to extract contractor name and address
        invoice.contractor_name = self._extract_contractor_name(lines)
        invoice.contractor_address = self._extract_address(lines)
        
        # Extract work items
        invoice.work_items = self._extract_work_items(text, lines)
        
        # Calculate total if not found
        if invoice.total == 0:
            invoice.calculate_total()
        
        return invoice
    
    def _extract_contractor_name(self, lines: List[str]) -> str:
        """Extract contractor name from document."""
        # Look for common patterns
        for i, line in enumerate(lines[:20]):  # Check first 20 lines
            line_lower = line.lower().strip()
            
            # Skip header lines
            if any(skip in line_lower for skip in ['invoice', 'date', 'from:', 'to:', 'bill to']):
                continue
            
            # Look for company indicators
            if any(indicator in line for indicator in ['Ltd', 'Limited', 'LLP', 'Inc', 'Company']):
                return line.strip()
            
            # Look for name patterns (2-3 words, capitalized)
            words = line.strip().split()
            if 2 <= len(words) <= 4:
                if all(w[0].isupper() for w in words if w and w[0].isalpha()):
                    return line.strip()
        
        # Fallback: return first non-empty line that's not a header
        for line in lines:
            stripped = line.strip()
            if stripped and len(stripped) > 3:
                if not any(skip in stripped.lower() for skip in ['invoice', 'date', 'page']):
                    return stripped
        
        return ""
    
    def _extract_address(self, lines: List[str]) -> str:
        """Extract address from document."""
        address_lines = []
        in_address = False
        
        for line in lines[:30]:
            stripped = line.strip()
            
            # Start of address section
            if any(marker in stripped.lower() for marker in ['address:', 'from:', 'contractor address']):
                in_address = True
                continue
            
            if in_address:
                # End of address (empty line or new section)
                if not stripped:
                    if address_lines:
                        break
                    continue
                
                # Check for address indicators
                if any(indicator in stripped for indicator in ['Street', 'Road', 'Lane', 'Avenue', 'Drive']):
                    address_lines.append(stripped)
                elif self.patterns["postcode"].search(stripped):
                    address_lines.append(stripped)
                    break
                elif len(stripped) > 5 and not any(skip in stripped.lower() for skip in ['invoice', 'date', 'email']):
                    address_lines.append(stripped)
        
        return ", ".join(address_lines) if address_lines else ""
    
    def _extract_work_items(self, text: str, lines: List[str]) -> List[WorkItem]:
        """Extract work items from document."""
        work_items = []
        
        # Look for table-like structures or itemized lists
        in_items_section = False
        
        for i, line in enumerate(lines):
            line_lower = line.lower().strip()
            
            # Detect start of items section
            if any(marker in line_lower for marker in ['description', 'items', 'services', 'work performed']):
                in_items_section = True
                continue
            
            if in_items_section:
                # Look for lines with amounts
                amount_match = self.patterns["amount"].search(line)
                if amount_match:
                    amount = float(amount_match.group(1).replace(',', ''))
                    
                    # Extract description (text before amount)
                    description = line[:amount_match.start()].strip()
                    if not description:
                        # Look at previous line for description
                        if i > 0:
                            description = lines[i-1].strip()
                    
                    if description and amount > 0:
                        work_item = WorkItem(
                            description=description,
                            amount=amount
                        )
                        work_items.append(work_item)
                
                # End of items section
                if any(marker in line_lower for marker in ['subtotal', 'total', 'vat', 'thank you']):
                    break
        
        return work_items


class MockDocumentProcessor(DocumentProcessor):
    """Mock processor for testing without OCR dependencies."""
    
    async def process_document(self, file_path: str, mime_type: str) -> InvoiceData:
        """Return sample data for testing."""
        logger.info(f"Mock processing: {file_path}")
        
        # Return sample invoice data
        invoice = InvoiceData(
            contractor_name="ABC Construction Ltd",
            contractor_address="123 Builder Street, London, EC1A 1BB",
            contractor_email="contact@abcconstruction.co.uk",
            contractor_utr="1234567890",
            contractor_ni="AB123456C",
            bank_account="12345678",
            sort_code="12-34-56",
            cardholder_name="John Smith",
            invoice_number="INV-2024-001",
            invoice_date="15/01/2024",
            work_start_date="01/01/2024",
            work_end_date="14/01/2024",
            operative_names="John Smith, Mike Johnson",
            subtotal=2500.00,
            vat_amount=500.00,
            cis_amount=500.00,
            source="upload",
        )
        
        invoice.work_items = [
            WorkItem(
                property_address="45 High Street, London",
                plot="A12",
                description="Kitchen renovation - plumbing and electrics",
                amount=1500.00,
            ),
            WorkItem(
                property_address="45 High Street, London",
                plot="A12",
                description="Bathroom tiling and fixtures",
                amount=1000.00,
            ),
        ]
        
        invoice.calculate_total()
        return invoice
