"""
Invoice Generator Module

Generates professional invoices in multiple formats (PDF, DOCX, HTML).
Uses Jinja2 templates for customizable invoice layouts.
"""

import os
import tempfile
from abc import ABC, abstractmethod
from dataclasses import dataclass
from datetime import datetime
from decimal import Decimal
from enum import Enum
from pathlib import Path
from typing import Optional, Dict, Any, List

from jinja2 import Environment, FileSystemLoader, Template

from src.core.config import get_settings
from src.models.invoice import InvoiceData, InvoiceItem, CompanyInfo
from src.utils.logger import get_logger
from src.utils.error_recovery import retry_with_backoff, ProcessingError

logger = get_logger(__name__)


class OutputFormat(Enum):
    """Supported output formats for invoices."""
    PDF = "pdf"
    DOCX = "docx"
    HTML = "html"


@dataclass
class GeneratedInvoice:
    """Result of invoice generation."""
    file_path: Path
    format: OutputFormat
    invoice_number: str
    file_size: int
    generated_at: datetime


class BaseInvoiceGenerator(ABC):
    """Abstract base class for invoice generators."""
    
    @abstractmethod
    async def generate(
        self, 
        invoice: InvoiceData, 
        output_path: Path
    ) -> Path:
        """Generate invoice file."""
        pass
    
    @abstractmethod
    def supports_format(self, format: OutputFormat) -> bool:
        """Check if generator supports the format."""
        pass


class HTMLInvoiceGenerator(BaseInvoiceGenerator):
    """Generate invoices as HTML."""
    
    DEFAULT_TEMPLATE = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Invoice {{ invoice.invoice_number }}</title>
    <style>
        * { margin: 0; padding: 0; box-sizing: border-box; }
        body {
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            line-height: 1.6;
            color: #333;
            padding: 40px;
            max-width: 800px;
            margin: 0 auto;
            background: #fff;
        }
        .invoice-header {
            display: flex;
            justify-content: space-between;
            margin-bottom: 40px;
            padding-bottom: 20px;
            border-bottom: 3px solid #2c3e50;
        }
        .company-info h1 {
            color: #2c3e50;
            font-size: 28px;
            margin-bottom: 10px;
        }
        .company-info p {
            color: #666;
            font-size: 14px;
        }
        .invoice-details {
            text-align: right;
        }
        .invoice-details h2 {
            color: #2c3e50;
            font-size: 32px;
            margin-bottom: 10px;
        }
        .invoice-details p {
            color: #666;
            font-size: 14px;
        }
        .bill-to {
            margin-bottom: 30px;
            padding: 20px;
            background: #f8f9fa;
            border-radius: 5px;
        }
        .bill-to h3 {
            color: #2c3e50;
            margin-bottom: 10px;
            font-size: 16px;
            text-transform: uppercase;
        }
        .bill-to p {
            font-size: 15px;
            line-height: 1.8;
        }
        .invoice-table {
            width: 100%;
            border-collapse: collapse;
            margin-bottom: 30px;
        }
        .invoice-table th {
            background: #2c3e50;
            color: white;
            padding: 12px;
            text-align: left;
            font-weight: 600;
        }
        .invoice-table td {
            padding: 12px;
            border-bottom: 1px solid #ddd;
        }
        .invoice-table tr:nth-child(even) {
            background: #f8f9fa;
        }
        .text-right { text-align: right; }
        .totals-section {
            margin-left: auto;
            width: 300px;
        }
        .totals-row {
            display: flex;
            justify-content: space-between;
            padding: 10px 0;
            border-bottom: 1px solid #ddd;
        }
        .totals-row.total {
            font-size: 20px;
            font-weight: bold;
            border-top: 3px solid #2c3e50;
            border-bottom: none;
            color: #2c3e50;
        }
        .notes-section {
            margin-top: 40px;
            padding-top: 20px;
            border-top: 1px solid #ddd;
        }
        .notes-section h4 {
            color: #2c3e50;
            margin-bottom: 10px;
        }
        .footer {
            margin-top: 50px;
            padding-top: 20px;
            border-top: 1px solid #ddd;
            text-align: center;
            color: #666;
            font-size: 12px;
        }
        @media print {
            body { padding: 20px; }
            .no-print { display: none; }
        }
    </style>
</head>
<body>
    <div class="invoice-header">
        <div class="company-info">
            <h1>{{ company.name or 'Your Company' }}</h1>
            {% if company.address %}<p>{{ company.address }}</p>{% endif %}
            {% if company.email %}<p>Email: {{ company.email }}</p>{% endif %}
            {% if company.phone %}<p>Phone: {{ company.phone }}</p>{% endif %}
            {% if company.tax_id %}<p>Tax ID: {{ company.tax_id }}</p>{% endif %}
        </div>
        <div class="invoice-details">
            <h2>INVOICE</h2>
            <p><strong>Invoice #:</strong> {{ invoice.invoice_number or 'DRAFT' }}</p>
            <p><strong>Issue Date:</strong> {{ invoice.issue_date }}</p>
            {% if invoice.due_date %}<p><strong>Due Date:</strong> {{ invoice.due_date }}</p>{% endif %}
            <p><strong>Status:</strong> {{ invoice.status.value|title }}</p>
        </div>
    </div>

    <div class="bill-to">
        <h3>Bill To:</h3>
        <p><strong>{{ invoice.customer.name }}</strong></p>
        {% if invoice.customer.address %}<p>{{ invoice.customer.address }}</p>{% endif %}
        {% if invoice.customer.email %}<p>Email: {{ invoice.customer.email }}</p>{% endif %}
        {% if invoice.customer.phone %}<p>Phone: {{ invoice.customer.phone }}</p>{% endif %}
        {% if invoice.customer.tax_id %}<p>Tax ID: {{ invoice.customer.tax_id }}</p>{% endif %}
    </div>

    <table class="invoice-table">
        <thead>
            <tr>
                <th>Description</th>
                <th class="text-right">Qty</th>
                <th class="text-right">Unit Price</th>
                {% if has_tax %}<th class="text-right">Tax</th>{% endif %}
                <th class="text-right">Amount</th>
            </tr>
        </thead>
        <tbody>
            {% for item in invoice.items %}
            <tr>
                <td>{{ item.description }}</td>
                <td class="text-right">{{ item.quantity }}</td>
                <td class="text-right">{{ currency_symbol }}{{ "%.2f"|format(item.unit_price|float) }}</td>
                {% if has_tax %}<td class="text-right">{{ item.tax_rate or 0 }}%</td>{% endif %}
                <td class="text-right">{{ currency_symbol }}{{ "%.2f"|format(item.amount|float) }}</td>
            </tr>
            {% endfor %}
        </tbody>
    </table>

    <div class="totals-section">
        <div class="totals-row">
            <span>Subtotal:</span>
            <span>{{ currency_symbol }}{{ "%.2f"|format(invoice.subtotal|float) }}</span>
        </div>
        {% if invoice.tax_total > 0 %}
        <div class="totals-row">
            <span>Tax:</span>
            <span>{{ currency_symbol }}{{ "%.2f"|format(invoice.tax_total|float) }}</span>
        </div>
        {% endif %}
        {% if invoice.discount > 0 %}
        <div class="totals-row">
            <span>Discount:</span>
            <span>-{{ currency_symbol }}{{ "%.2f"|format(invoice.discount|float) }}</span>
        </div>
        {% endif %}
        <div class="totals-row total">
            <span>Total:</span>
            <span>{{ currency_symbol }}{{ "%.2f"|format(invoice.total|float) }}</span>
        </div>
    </div>

    {% if invoice.notes %}
    <div class="notes-section">
        <h4>Notes:</h4>
        <p>{{ invoice.notes }}</p>
    </div>
    {% endif %}

    {% if invoice.terms %}
    <div class="notes-section">
        <h4>Terms & Conditions:</h4>
        <p>{{ invoice.terms }}</p>
    </div>
    {% endif %}

    <div class="footer">
        <p>Thank you for your business!</p>
        <p>Payment Terms: {{ invoice.payment_terms.value }}</p>
    </div>
</body>
</html>
"""
    
    CURRENCY_SYMBOLS = {
        "USD": "$",
        "EUR": "€",
        "GBP": "£",
        "JPY": "¥",
        "CAD": "C$",
        "AUD": "A$",
    }
    
    def __init__(self):
        self.settings = get_settings()
        self.template_dir = Path(__file__).parent.parent / "templates"
        self._env = None
    
    def _get_jinja_env(self) -> Environment:
        """Get or create Jinja environment."""
        if self._env is None:
            if self.template_dir.exists():
                self._env = Environment(loader=FileSystemLoader(self.template_dir))
            else:
                self._env = Environment(loader=FileSystemLoader("/"))
        return self._env
    
    def supports_format(self, format: OutputFormat) -> bool:
        """Check if HTML generator supports the format."""
        return format == OutputFormat.HTML
    
    @retry_with_backoff(max_attempts=3)
    async def generate(
        self, 
        invoice: InvoiceData, 
        output_path: Path
    ) -> Path:
        """Generate HTML invoice."""
        try:
            template = self._get_template()
            
            # Prepare template context
            context = self._prepare_context(invoice)
            
            # Render HTML
            html_content = template.render(**context)
            
            # Write to file
            output_path = Path(output_path)
            output_path.write_text(html_content, encoding="utf-8")
            
            logger.info(f"Generated HTML invoice: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"HTML generation failed: {e}")
            raise ProcessingError(f"Failed to generate HTML invoice: {e}")
    
    def _get_template(self) -> Template:
        """Get the invoice template."""
        env = self._get_jinja_env()
        
        # Try to load custom template
        template_path = self.template_dir / "invoice.html"
        if template_path.exists():
            return env.get_template("invoice.html")
        
        # Use default template
        return env.from_string(self.DEFAULT_TEMPLATE)
    
    def _prepare_context(self, invoice: InvoiceData) -> Dict[str, Any]:
        """Prepare template context."""
        # Get company info from settings or invoice
        company = invoice.company or CompanyInfo(
            name=self.settings.company_name,
            address=self.settings.company_address,
            email=self.settings.company_email,
            phone=self.settings.company_phone,
            tax_id=self.settings.company_tax_id,
        )
        
        # Check if any items have tax
        has_tax = any(item.tax_rate for item in invoice.items)
        
        # Get currency symbol
        currency_symbol = self.CURRENCY_SYMBOLS.get(invoice.currency, invoice.currency)
        
        return {
            "invoice": invoice,
            "company": company,
            "has_tax": has_tax,
            "currency_symbol": currency_symbol,
        }


class PDFInvoiceGenerator(BaseInvoiceGenerator):
    """Generate invoices as PDF using WeasyPrint."""
    
    def __init__(self):
        self.settings = get_settings()
        self.html_generator = HTMLInvoiceGenerator()
    
    def supports_format(self, format: OutputFormat) -> bool:
        """Check if PDF generator supports the format."""
        return format == OutputFormat.PDF
    
    @retry_with_backoff(max_attempts=3)
    async def generate(
        self, 
        invoice: InvoiceData, 
        output_path: Path
    ) -> Path:
        """Generate PDF invoice from HTML."""
        try:
            from weasyprint import HTML, CSS
            
            # First generate HTML
            with tempfile.NamedTemporaryFile(
                mode="w", 
                suffix=".html", 
                delete=False
            ) as tmp_html:
                template = self.html_generator._get_template()
                context = self.html_generator._prepare_context(invoice)
                html_content = template.render(**context)
                tmp_html.write(html_content)
                tmp_html_path = tmp_html.name
            
            try:
                # Convert to PDF
                output_path = Path(output_path)
                HTML(filename=tmp_html_path).write_pdf(str(output_path))
                
                logger.info(f"Generated PDF invoice: {output_path}")
                return output_path
                
            finally:
                # Cleanup temp file
                Path(tmp_html_path).unlink(missing_ok=True)
                
        except ImportError:
            logger.error("weasyprint not installed")
            raise ProcessingError("PDF generation requires weasyprint")
        except Exception as e:
            logger.error(f"PDF generation failed: {e}")
            raise ProcessingError(f"Failed to generate PDF invoice: {e}")


class DOCXInvoiceGenerator(BaseInvoiceGenerator):
    """Generate invoices as Word documents."""
    
    def __init__(self):
        self.settings = get_settings()
    
    def supports_format(self, format: OutputFormat) -> bool:
        """Check if DOCX generator supports the format."""
        return format == OutputFormat.DOCX
    
    @retry_with_backoff(max_attempts=3)
    async def generate(
        self, 
        invoice: InvoiceData, 
        output_path: Path
    ) -> Path:
        """Generate DOCX invoice."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # Get company info
            company = invoice.company or CompanyInfo(
                name=self.settings.company_name,
                address=self.settings.company_address,
                email=self.settings.company_email,
                phone=self.settings.company_phone,
                tax_id=self.settings.company_tax_id,
            )
            
            # Header - Company Info
            header = doc.add_heading(company.name or "Your Company", 0)
            
            if company.address:
                doc.add_paragraph(company.address)
            if company.email:
                doc.add_paragraph(f"Email: {company.email}")
            if company.phone:
                doc.add_paragraph(f"Phone: {company.phone}")
            
            doc.add_paragraph()  # Spacing
            
            # Invoice Details
            details = doc.add_paragraph()
            details.add_run("INVOICE\n").bold = True
            details.add_run(f"Invoice #: {invoice.invoice_number or 'DRAFT'}\n")
            details.add_run(f"Issue Date: {invoice.issue_date}\n")
            if invoice.due_date:
                details.add_run(f"Due Date: {invoice.due_date}\n")
            details.add_run(f"Status: {invoice.status.value.title()}")
            
            doc.add_paragraph()  # Spacing
            
            # Bill To
            doc.add_heading("Bill To:", level=2)
            doc.add_paragraph(invoice.customer.name).bold = True
            if invoice.customer.address:
                doc.add_paragraph(invoice.customer.address)
            if invoice.customer.email:
                doc.add_paragraph(f"Email: {invoice.customer.email}")
            if invoice.customer.phone:
                doc.add_paragraph(f"Phone: {invoice.customer.phone}")
            
            doc.add_paragraph()  # Spacing
            
            # Items Table
            table = doc.add_table(rows=1, cols=4)
            table.style = "Table Grid"
            
            # Header row
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Description"
            hdr_cells[1].text = "Qty"
            hdr_cells[2].text = "Unit Price"
            hdr_cells[3].text = "Amount"
            
            # Make header bold
            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            # Add items
            for item in invoice.items:
                row_cells = table.add_row().cells
                row_cells[0].text = item.description
                row_cells[1].text = str(item.quantity)
                row_cells[2].text = f"{invoice.currency} {item.unit_price:.2f}"
                row_cells[3].text = f"{invoice.currency} {item.amount:.2f}"
            
            doc.add_paragraph()  # Spacing
            
            # Totals
            totals = doc.add_paragraph()
            totals.add_run(f"Subtotal: {invoice.currency} {invoice.subtotal:.2f}\n")
            if invoice.tax_total > 0:
                totals.add_run(f"Tax: {invoice.currency} {invoice.tax_total:.2f}\n")
            if invoice.discount > 0:
                totals.add_run(f"Discount: -{invoice.currency} {invoice.discount:.2f}\n")
            total_run = totals.add_run(f"Total: {invoice.currency} {invoice.total:.2f}")
            total_run.bold = True
            total_run.font.size = Pt(14)
            
            # Notes
            if invoice.notes:
                doc.add_paragraph()
                doc.add_heading("Notes:", level=2)
                doc.add_paragraph(invoice.notes)
            
            # Terms
            if invoice.terms:
                doc.add_paragraph()
                doc.add_heading("Terms & Conditions:", level=2)
                doc.add_paragraph(invoice.terms)
            
            # Save
            output_path = Path(output_path)
            doc.save(str(output_path))
            
            logger.info(f"Generated DOCX invoice: {output_path}")
            return output_path
            
        except ImportError:
            logger.error("python-docx not installed")
            raise ProcessingError("DOCX generation requires python-docx")
        except Exception as e:
            logger.error(f"DOCX generation failed: {e}")
            raise ProcessingError(f"Failed to generate DOCX invoice: {e}")


class InvoiceGeneratorService:
    """
    Service for generating invoices in multiple formats.
    
    Coordinates different generators and provides a unified interface.
    """
    
    def __init__(self):
        self.settings = get_settings()
        self.generators: List[BaseInvoiceGenerator] = [
            HTMLInvoiceGenerator(),
            PDFInvoiceGenerator(),
            DOCXInvoiceGenerator(),
        ]
    
    def get_supported_formats(self) -> List[OutputFormat]:
        """Get list of supported output formats."""
        return [f for f in OutputFormat]
    
    def get_file_extension(self, format: OutputFormat) -> str:
        """Get file extension for format."""
        return format.value
    
    async def generate_invoice(
        self,
        invoice: InvoiceData,
        format: OutputFormat = OutputFormat.PDF,
        output_dir: Optional[Path] = None,
        filename: Optional[str] = None,
    ) -> GeneratedInvoice:
        """
        Generate an invoice in the specified format.
        
        Args:
            invoice: Invoice data to generate
            format: Output format
            output_dir: Optional output directory
            filename: Optional filename (without extension)
            
        Returns:
            GeneratedInvoice: Result with file path
        """
        # Determine output directory
        if output_dir is None:
            output_dir = self.settings.invoice_output_dir
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)
        
        # Generate filename
        if filename is None:
            timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
            invoice_num = invoice.invoice_number or f"DRAFT_{timestamp}"
            filename = f"{invoice_num}"
        
        output_path = output_dir / f"{filename}.{format.value}"
        
        # Find appropriate generator
        generator = None
        for g in self.generators:
            if g.supports_format(format):
                generator = g
                break
        
        if not generator:
            raise ProcessingError(f"No generator available for format: {format.value}")
        
        # Generate invoice
        file_path = await generator.generate(invoice, output_path)
        
        return GeneratedInvoice(
            file_path=file_path,
            format=format,
            invoice_number=invoice.invoice_number or "DRAFT",
            file_size=file_path.stat().st_size,
            generated_at=datetime.utcnow(),
        )
    
    async def generate_multiple(
        self,
        invoice: InvoiceData,
        formats: List[OutputFormat],
        output_dir: Optional[Path] = None,
    ) -> List[GeneratedInvoice]:
        """
        Generate invoice in multiple formats.
        
        Args:
            invoice: Invoice data
            formats: List of formats to generate
            output_dir: Optional output directory
            
        Returns:
            List[GeneratedInvoice]: Results for each format
        """
        results = []
        for fmt in formats:
            try:
                result = await self.generate_invoice(invoice, fmt, output_dir)
                results.append(result)
            except Exception as e:
                logger.error(f"Failed to generate {fmt.value}: {e}")
        return results


# Global service instance
_service: Optional[InvoiceGeneratorService] = None


def get_invoice_generator() -> InvoiceGeneratorService:
    """
    Get the global invoice generator service.
    
    Returns:
        InvoiceGeneratorService: Singleton instance
    """
    global _service
    if _service is None:
        _service = InvoiceGeneratorService()
    return _service
