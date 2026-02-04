"""
Document Processor Module

Handles multi-format document processing including PDF, images (JPEG, PNG),
and DOCX files. Extracts text content for AI analysis.
"""

import os
import io
import tempfile
from abc import ABC, abstractmethod
from dataclasses import dataclass
from enum import Enum
from pathlib import Path
from typing import Optional, Dict, Any, List, BinaryIO
from datetime import datetime

import aiofiles
from PIL import Image

from src.core.config import get_settings
from src.utils.logger import get_logger
from src.utils.error_recovery import retry_with_backoff, ProcessingError

logger = get_logger(__name__)


class DocumentType(Enum):
    """Supported document types."""
    PDF = "pdf"
    JPEG = "jpeg"
    JPG = "jpg"
    PNG = "png"
    DOCX = "docx"
    TEXT = "txt"
    UNKNOWN = "unknown"


@dataclass
class ProcessedDocument:
    """Result of document processing."""
    file_path: Path
    file_name: str
    file_type: DocumentType
    file_size: int
    extracted_text: str
    page_count: Optional[int] = None
    image_count: Optional[int] = None
    processing_time_ms: Optional[float] = None
    metadata: Dict[str, Any] = None
    
    def __post_init__(self):
        if self.metadata is None:
            self.metadata = {}


class BaseDocumentHandler(ABC):
    """Abstract base class for document handlers."""
    
    @abstractmethod
    async def can_handle(self, file_path: Path) -> bool:
        """Check if this handler can process the file."""
        pass
    
    @abstractmethod
    async def extract_text(self, file_path: Path) -> str:
        """Extract text content from the file."""
        pass
    
    @abstractmethod
    async def get_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract metadata from the file."""
        pass


class PDFHandler(BaseDocumentHandler):
    """Handler for PDF documents."""
    
    async def can_handle(self, file_path: Path) -> bool:
        """Check if file is a PDF."""
        return file_path.suffix.lower() == ".pdf"
    
    @retry_with_backoff(max_attempts=3)
    async def extract_text(self, file_path: Path) -> str:
        """Extract text from PDF using PyPDF2."""
        try:
            import PyPDF2
            
            text_parts = []
            page_count = 0
            
            with open(file_path, "rb") as file:
                reader = PyPDF2.PdfReader(file)
                page_count = len(reader.pages)
                
                for page_num, page in enumerate(reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_parts.append(f"--- Page {page_num + 1} ---\n{page_text}")
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
            
            extracted_text = "\n\n".join(text_parts)
            
            if not extracted_text.strip():
                # Try OCR for scanned PDFs
                logger.info("No text extracted, attempting OCR on PDF")
                extracted_text = await self._extract_with_ocr(file_path)
            
            logger.info(f"Extracted {len(extracted_text)} characters from {page_count} pages")
            return extracted_text
            
        except Exception as e:
            logger.error(f"PDF text extraction failed: {e}")
            raise ProcessingError(f"Failed to extract text from PDF: {e}")
    
    async def _extract_with_ocr(self, file_path: Path) -> str:
        """Extract text from scanned PDF using OCR."""
        try:
            from pdf2image import convert_from_path
            import pytesseract
            
            images = convert_from_path(str(file_path))
            text_parts = []
            
            for i, image in enumerate(images):
                text = pytesseract.image_to_string(image)
                text_parts.append(f"--- Page {i + 1} ---\n{text}")
            
            return "\n\n".join(text_parts)
            
        except ImportError:
            logger.warning("OCR dependencies not available (pdf2image, pytesseract)")
            return ""
        except Exception as e:
            logger.error(f"OCR extraction failed: {e}")
            return ""
    
    async def get_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract PDF metadata."""
        try:
            import PyPDF2
            
            with open(file_path, "rb") as file:
                reader = PyPDF2.PdfReader(file)
                info = reader.metadata
                
                return {
                    "page_count": len(reader.pages),
                    "author": info.author if info else None,
                    "creator": info.creator if info else None,
                    "producer": info.producer if info else None,
                    "subject": info.subject if info else None,
                    "title": info.title if info else None,
                }
        except Exception as e:
            logger.warning(f"Failed to extract PDF metadata: {e}")
            return {"page_count": 0}


class ImageHandler(BaseDocumentHandler):
    """Handler for image documents (JPEG, PNG)."""
    
    SUPPORTED_FORMATS = {".jpg", ".jpeg", ".png"}
    
    async def can_handle(self, file_path: Path) -> bool:
        """Check if file is an image."""
        return file_path.suffix.lower() in self.SUPPORTED_FORMATS
    
    @retry_with_backoff(max_attempts=3)
    async def extract_text(self, file_path: Path) -> str:
        """Extract text from image using OCR."""
        try:
            import pytesseract
            
            # Open and preprocess image
            with Image.open(file_path) as img:
                # Convert to RGB if necessary
                if img.mode != "RGB":
                    img = img.convert("RGB")
                
                # Enhance image for better OCR
                img = self._preprocess_image(img)
                
                # Extract text
                text = pytesseract.image_to_string(img)
                
                logger.info(f"Extracted {len(text)} characters from image")
                return text
                
        except ImportError:
            logger.error("pytesseract not installed")
            raise ProcessingError("OCR not available. Install pytesseract.")
        except Exception as e:
            logger.error(f"Image text extraction failed: {e}")
            raise ProcessingError(f"Failed to extract text from image: {e}")
    
    def _preprocess_image(self, img: Image.Image) -> Image.Image:
        """Preprocess image for better OCR results."""
        # Resize if too large (maintain aspect ratio)
        max_dimension = 3000
        if max(img.size) > max_dimension:
            ratio = max_dimension / max(img.size)
            new_size = (int(img.size[0] * ratio), int(img.size[1] * ratio))
            img = img.resize(new_size, Image.Resampling.LANCZOS)
        
        return img
    
    async def get_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract image metadata."""
        try:
            with Image.open(file_path) as img:
                return {
                    "format": img.format,
                    "mode": img.mode,
                    "width": img.width,
                    "height": img.height,
                    "size": img.size,
                }
        except Exception as e:
            logger.warning(f"Failed to extract image metadata: {e}")
            return {}


class DOCXHandler(BaseDocumentHandler):
    """Handler for Word documents."""
    
    async def can_handle(self, file_path: Path) -> bool:
        """Check if file is a DOCX."""
        return file_path.suffix.lower() == ".docx"
    
    @retry_with_backoff(max_attempts=3)
    async def extract_text(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        try:
            from docx import Document
            
            doc = Document(file_path)
            text_parts = []
            
            # Extract from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            extracted_text = "\n".join(text_parts)
            logger.info(f"Extracted {len(extracted_text)} characters from DOCX")
            return extracted_text
            
        except Exception as e:
            logger.error(f"DOCX text extraction failed: {e}")
            raise ProcessingError(f"Failed to extract text from DOCX: {e}")
    
    async def get_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract DOCX metadata."""
        try:
            from docx import Document
            
            doc = Document(file_path)
            core_props = doc.core_properties
            
            return {
                "author": core_props.author,
                "title": core_props.title,
                "subject": core_props.subject,
                "created": core_props.created.isoformat() if core_props.created else None,
                "modified": core_props.modified.isoformat() if core_props.modified else None,
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
            }
        except Exception as e:
            logger.warning(f"Failed to extract DOCX metadata: {e}")
            return {}


class DocumentProcessor:
    """
    Main document processor that coordinates different handlers.
    
    Provides a unified interface for processing various document formats
    and extracting text content.
    """
    
    def __init__(self):
        """Initialize the document processor with handlers."""
        self.settings = get_settings()
        self.handlers: List[BaseDocumentHandler] = [
            PDFHandler(),
            ImageHandler(),
            DOCXHandler(),
        ]
    
    def get_document_type(self, file_path: Path) -> DocumentType:
        """Determine document type from file extension."""
        ext = file_path.suffix.lower()
        
        type_map = {
            ".pdf": DocumentType.PDF,
            ".jpg": DocumentType.JPG,
            ".jpeg": DocumentType.JPEG,
            ".png": DocumentType.PNG,
            ".docx": DocumentType.DOCX,
            ".txt": DocumentType.TEXT,
        }
        
        return type_map.get(ext, DocumentType.UNKNOWN)
    
    def is_supported(self, file_path: Path) -> bool:
        """Check if file format is supported."""
        doc_type = self.get_document_type(file_path)
        return doc_type != DocumentType.UNKNOWN
    
    async def process_file(self, file_path: Path) -> ProcessedDocument:
        """
        Process a document file and extract text.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            ProcessedDocument: Extracted content and metadata
            
        Raises:
            ProcessingError: If processing fails
            ValueError: If file format not supported
        """
        import time
        
        start_time = time.time()
        file_path = Path(file_path)
        
        # Validate file exists
        if not file_path.exists():
            raise ProcessingError(f"File not found: {file_path}")
        
        # Check file size
        file_size = file_path.stat().st_size
        if file_size > self.settings.max_file_size_bytes:
            raise ProcessingError(
                f"File too large: {file_size} bytes (max: {self.settings.max_file_size_bytes})"
            )
        
        # Determine document type
        doc_type = self.get_document_type(file_path)
        if doc_type == DocumentType.UNKNOWN:
            raise ValueError(f"Unsupported file format: {file_path.suffix}")
        
        # Find appropriate handler
        handler = None
        for h in self.handlers:
            if await h.can_handle(file_path):
                handler = h
                break
        
        if not handler:
            raise ProcessingError(f"No handler available for {doc_type.value}")
        
        # Extract text and metadata
        try:
            extracted_text = await handler.extract_text(file_path)
            metadata = await handler.get_metadata(file_path)
            
            processing_time = (time.time() - start_time) * 1000
            
            return ProcessedDocument(
                file_path=file_path,
                file_name=file_path.name,
                file_type=doc_type,
                file_size=file_size,
                extracted_text=extracted_text,
                page_count=metadata.get("page_count"),
                image_count=metadata.get("image_count"),
                processing_time_ms=processing_time,
                metadata=metadata,
            )
            
        except Exception as e:
            logger.error(f"Document processing failed: {e}")
            raise ProcessingError(f"Failed to process document: {e}")
    
    async def process_bytes(
        self, 
        data: bytes, 
        file_name: str,
        file_type: Optional[str] = None
    ) -> ProcessedDocument:
        """
        Process document from bytes.
        
        Args:
            data: File content as bytes
            file_name: Original file name
            file_type: Optional MIME type
            
        Returns:
            ProcessedDocument: Extracted content and metadata
        """
        # Save to temporary file
        suffix = Path(file_name).suffix
        
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(data)
            tmp_path = Path(tmp.name)
        
        try:
            result = await self.process_file(tmp_path)
            # Update with original filename
            result.file_name = file_name
            return result
        finally:
            # Cleanup temporary file
            try:
                tmp_path.unlink()
            except Exception:
                pass
    
    async def save_upload(
        self, 
        data: bytes, 
        file_name: str,
        user_id: int
    ) -> Path:
        """
        Save uploaded file to storage.
        
        Args:
            data: File content
            file_name: Original file name
            user_id: Uploading user ID
            
        Returns:
            Path: Path to saved file
        """
        # Create user directory
        user_dir = self.settings.upload_dir / str(user_id)
        user_dir.mkdir(parents=True, exist_ok=True)
        
        # Generate unique filename
        timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        safe_name = "".join(c for c in file_name if c.isalnum() or c in "._-")
        unique_name = f"{timestamp}_{safe_name}"
        
        file_path = user_dir / unique_name
        
        async with aiofiles.open(file_path, "wb") as f:
            await f.write(data)
        
        logger.info(f"Saved upload to {file_path}")
        return file_path


# Global processor instance
_processor: Optional[DocumentProcessor] = None


def get_document_processor() -> DocumentProcessor:
    """
    Get the global document processor instance.
    
    Returns:
        DocumentProcessor: Singleton instance
    """
    global _processor
    if _processor is None:
        _processor = DocumentProcessor()
    return _processor
