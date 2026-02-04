"""
Invoice Processing Pipeline - DOCX Extractor
Handles text and table extraction from Word documents.
"""

import os
import logging
from typing import Optional, List, Dict, Any
from pathlib import Path
import re

from ..models import InvoiceData, ExtractionResult, Contractor, Address, BankDetails, Financials, WorkItem, WorkPeriod
from ..utils import (
    clean_text, InvoicePatterns, extract_pattern, parse_date_flexible,
    parse_money, extract_address, parse_date_range
)

logger = logging.getLogger(__name__)

# Try to import python-docx
try:
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available")


class DOCXTextExtractor:
    """Extract text from DOCX files."""
    
    def __init__(self):
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx is required for DOCX extraction")
    
    def extract_text(self, file_path: str) -> str:
        """
        Extract all text from DOCX file.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text
        """
        doc = Document(file_path)
        
        text_parts = []
        
        # Extract from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    text_parts.append(' | '.join(row_text))
        
        return '\n'.join(text_parts)
    
    def extract_paragraphs(self, file_path: str) -> List[str]:
        """Extract text from paragraphs only."""
        doc = Document(file_path)
        return [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    
    def extract_tables(self, file_path: str) -> List[List[List[str]]]:
        """
        Extract all tables from DOCX.
        
        Returns:
            List of tables, each table is a list of rows
        """
        doc = Document(file_path)
        tables = []
        
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
        
        return tables


class DOCXInvoiceExtractor:
    """Extract invoice data from DOCX files."""
    
    def __init__(self):
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx is required for DOCX extraction")
        self.text_extractor = DOCXTextExtractor()
    
    def extract(self, file_path: str) -> ExtractionResult:
        """
        Extract invoice data from DOCX file.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            ExtractionResult with invoice data
        """
        import time
        start_time = time.time()
        
        try:
            # Verify file exists
            if not os.path.exists(file_path):
                return ExtractionResult(
                    success=False,
                    method='docx_extraction',
                    error_message=f"File not found: {file_path}"
                )
            
            # Extract text
            text = self.text_extractor.extract_text(file_path)
            
            if not text.strip():
                return ExtractionResult(
                    success=False,
                    method='docx_extraction',
                    error_message="No text extracted from DOCX"
                )
            
            # Parse invoice data
            invoice_data = self._parse_invoice_text(text)
            invoice_data.raw_text = text[:5000]
            
            # Also extract tables
            try:
                tables = self.text_extractor.extract_tables(file_path)
                if tables:
                    for table in tables:
                        self._merge_table_data(invoice_data, table)
            except Exception as e:
                logger.warning(f"Table extraction failed: {e}")
            
            processing_time = int((time.time() - start_time) * 1000)
            
            return ExtractionResult(
                success=True,
                method='docx_extraction',
                data=invoice_data,
                confidence=invoice_data.extraction_confidence,
                processing_time_ms=processing_time,
                raw_text=text[:2000]
            )
            
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            processing_time = int((time.time() - start_time) * 1000)
            
            return ExtractionResult(
                success=False,
                method='docx_extraction',
                error_message=str(e),
                processing_time_ms=processing_time
            )
    
    def _parse_invoice_text(self, text: str) -> InvoiceData:
        """Parse invoice data from extracted text."""
        invoice_data = InvoiceData(extraction_method='docx_extraction')
        
        # Extract invoice number
        invoice_number = extract_pattern(text, InvoicePatterns.INVOICE_NUMBER)
        if invoice_number:
            invoice_data.invoice_number = invoice_number
        
        # Extract invoice date
        for pattern in InvoicePatterns.DATE_PATTERNS:
            date_str = extract_pattern(text, pattern)
            if date_str:
                invoice_data.invoice_date = parse_date_flexible(date_str)
                if invoice_data.invoice_date:
                    break
        
        # Extract work period
        period_start = extract_pattern(text, InvoicePatterns.PERIOD_START)
        period_end = extract_pattern(text, InvoicePatterns.PERIOD_END)
        
        if period_start or period_end:
            invoice_data.work_period = WorkPeriod(
                start_date=parse_date_flexible(period_start) if period_start else None,
                end_date=parse_date_flexible(period_end) if period_end else None
            )
        else:
            start, end = parse_date_range(text)
            if start or end:
                invoice_data.work_period = WorkPeriod(start_date=start, end_date=end)
        
        # Extract contractor information
        contractor = Contractor()
        
        utr = extract_pattern(text, InvoicePatterns.UTR)
        if utr:
            contractor.utr = utr
        
        ni = extract_pattern(text, InvoicePatterns.NI_NUMBER)
        if ni:
            contractor.ni_number = ni
        
        vat = extract_pattern(text, InvoicePatterns.VAT_NUMBER)
        if vat:
            contractor.vat_number = vat
        
        company = extract_pattern(text, InvoicePatterns.COMPANY_NUMBER)
        if company:
            contractor.company_number = company
        
        email = extract_pattern(text, InvoicePatterns.EMAIL)
        if email:
            contractor.email = email.lower()
        
        phone = extract_pattern(text, InvoicePatterns.PHONE)
        if phone:
            contractor.phone = phone
        
        bank = BankDetails()
        sort_code = extract_pattern(text, InvoicePatterns.SORT_CODE)
        if sort_code:
            bank.sort_code = sort_code
        
        account = extract_pattern(text, InvoicePatterns.ACCOUNT_NUMBER)
        if account:
            bank.account_number = account
        
        if sort_code or account:
            contractor.bank_details = bank
        
        address_info = extract_address(text)
        if any(address_info.values()):
            contractor.address = Address(**address_info)
        
        # Try to find contractor name
        contractor_patterns = [
            r'(?:From|Contractor|Subcontractor|Supplier)[:\s]*\n?\s*([A-Z][A-Za-z0-9\s&.,]+(?:Ltd|Limited|LLP|Inc|Co|Company)?)',
            r'^([A-Z][A-Za-z\s&.,]+(?:Ltd|Limited|LLP|Inc|Co|Company)?)',
        ]
        
        for pattern in contractor_patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
            if match:
                contractor.name = match.group(1).strip()
                break
        
        if any([contractor.name, contractor.utr, contractor.email]):
            invoice_data.contractor = contractor
        
        # Extract financials
        financials = Financials()
        
        subtotal = extract_pattern(text, InvoicePatterns.SUBTOTAL)
        if subtotal:
            financials.subtotal = parse_money(subtotal)
        
        vat_amount = extract_pattern(text, InvoicePatterns.VAT_AMOUNT)
        if vat_amount:
            financials.vat_amount = parse_money(vat_amount)
        
        vat_rate = extract_pattern(text, InvoicePatterns.VAT_RATE)
        if vat_rate:
            try:
                financials.vat_rate = parse_money(vat_rate)
            except:
                pass
        
        cis = extract_pattern(text, InvoicePatterns.CIS_DEDUCTION)
        if cis:
            financials.cis_deduction = parse_money(cis)
        
        total = extract_pattern(text, InvoicePatterns.TOTAL_DUE)
        if total:
            financials.total_due = parse_money(total)
        
        if any([financials.subtotal, financials.total_due]):
            invoice_data.financials = financials
        
        # Extract work items
        work_items = self._extract_work_items(text)
        if work_items:
            invoice_data.work_items = work_items
        
        # Calculate confidence
        invoice_data.extraction_confidence = invoice_data.completeness_score()
        
        return invoice_data
    
    def _merge_table_data(self, invoice_data: InvoiceData, table: List[List[str]]):
        """Merge data from extracted table into invoice data."""
        if not table or len(table) < 2:
            return
        
        # Try to identify columns
        headers = [str(h).lower().strip() if h else '' for h in table[0]]
        
        # Find column indices
        desc_idx = next((i for i, h in enumerate(headers) if 'desc' in h or 'item' in h), None)
        qty_idx = next((i for i, h in enumerate(headers) if 'qty' in h or 'quantity' in h), None)
        price_idx = next((i for i, h in enumerate(headers) if 'price' in h or 'rate' in h), None)
        amount_idx = next((i for i, h in enumerate(headers) if 'amount' in h or 'total' in h), None)
        plot_idx = next((i for i, h in enumerate(headers) if 'plot' in h or 'property' in h), None)
        
        # Process data rows
        for row in table[1:]:
            if not any(row):
                continue
            
            work_item = WorkItem()
            
            if plot_idx is not None and plot_idx < len(row):
                work_item.plot_number = str(row[plot_idx]).strip()
            
            if desc_idx is not None and desc_idx < len(row):
                work_item.description = str(row[desc_idx]).strip()
            
            if qty_idx is not None and qty_idx < len(row):
                try:
                    work_item.quantity = parse_money(str(row[qty_idx]))
                except:
                    pass
            
            if price_idx is not None and price_idx < len(row):
                try:
                    work_item.unit_price = parse_money(str(row[price_idx]))
                except:
                    pass
            
            if amount_idx is not None and amount_idx < len(row):
                try:
                    work_item.amount = parse_money(str(row[amount_idx]))
                except:
                    pass
            
            if any([work_item.description, work_item.plot_number, work_item.amount]):
                invoice_data.work_items.append(work_item)
    
    def _extract_work_items(self, text: str) -> List[WorkItem]:
        """Extract work items from text."""
        work_items = []
        
        # Look for plot/property references
        plots = InvoicePatterns.PLOT_NUMBER.findall(text)
        
        for plot in plots:
            item = WorkItem(plot_number=plot)
            
            # Try to find associated description
            pattern = re.compile(
                rf'(?:Plot|Property)[:\s#]*{re.escape(plot)}.*?\n(.{{50,200}})',
                re.IGNORECASE | re.DOTALL
            )
            match = pattern.search(text)
            if match:
                item.description = match.group(1).strip()
            
            work_items.append(item)
        
        # If no plots found, look for operative names
        if not work_items:
            operatives = InvoicePatterns.OPERATIVE.findall(text)
            for op in operatives:
                item = WorkItem(operative_names=[op])
                work_items.append(item)
        
        return work_items


def extract_from_docx(file_path: str) -> ExtractionResult:
    """Convenience function for DOCX extraction."""
    extractor = DOCXInvoiceExtractor()
    return extractor.extract(file_path)
