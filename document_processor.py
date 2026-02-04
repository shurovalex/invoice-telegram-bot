"""
Document Processor Module for Invoice Agent
Handles PDF, JPEG, PNG, and DOCX files with multiple extraction strategies.
"""

import re
import logging
from typing import Dict, List, Optional, Any, Union
from dataclasses import dataclass, field
from pathlib import Path

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False
try:
    import pytesseract
    from PIL import Image
    TESSERACT_AVAILABLE = True
except ImportError:
    TESSERACT_AVAILABLE = False
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False


@dataclass
class ExtractedInvoice:
    """Structured invoice data container."""
    vendor_name: Optional[str] = None
    invoice_number: Optional[str] = None
    invoice_date: Optional[str] = None
    due_date: Optional[str] = None
    total_amount: Optional[float] = None
    subtotal: Optional[float] = None
    tax_amount: Optional[float] = None
    currency: Optional[str] = None
    line_items: List[Dict[str, Any]] = field(default_factory=list)
    raw_text: str = ""
    confidence: float = 0.0
    source: str = ""
    errors: List[str] = field(default_factory=list)

    def to_dict(self) -> Dict[str, Any]:
        return {
            "vendor_name": self.vendor_name, "invoice_number": self.invoice_number,
            "invoice_date": self.invoice_date, "due_date": self.due_date,
            "total_amount": self.total_amount, "subtotal": self.subtotal,
            "tax_amount": self.tax_amount, "currency": self.currency,
            "line_items": self.line_items, "raw_text": self.raw_text[:5000],
            "confidence": self.confidence, "source": self.source, "errors": self.errors
        }


class DocumentProcessor:
    """Process documents to extract invoice data with multiple strategies."""
    
    PATTERNS = {
        "invoice_number": [
            r'(?:invoice|inv|bill)\s*#?\s*:?\s*([A-Z0-9\-]{3,20})',
            r'(?:invoice|inv)\s+(?:number|no|#)\s*:?\s*([A-Z0-9\-]{3,20})',
            r'#\s*([A-Z0-9\-]{5,15})',
        ],
        "date": [r'(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})', r'(\d{4}[/-]\d{1,2}[/-]\d{1,2})'],
        "amount": [
            r'(?:total|amount|due)\s*:?\s*[$€£¥]?\s*([\d,]+\.\d{2})',
            r'[$€£¥]\s*([\d,]+\.\d{2})',
            r'balance\s+due\s*:?\s*[$€£¥]?\s*([\d,]+\.\d{2})',
        ],
        "currency": [r'[$]\s*\d', r'USD|EUR|GBP|JPY|CAD|AUD']
    }
    
    def __init__(self, tesseract_cmd: Optional[str] = None):
        if tesseract_cmd and TESSERACT_AVAILABLE:
            pytesseract.pytesseract.tesseract_cmd = tesseract_cmd
    
    def process(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """Process any supported document type."""
        path = Path(file_path)
        ext = path.suffix.lower()
        try:
            if ext == '.pdf': return self.process_pdf(path)
            elif ext in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff']: return self.process_image(path)
            elif ext == '.docx': return self.process_docx(path)
            else: return ExtractedInvoice(errors=[f"Unsupported: {ext}"]).to_dict()
        except Exception as e:
            return ExtractedInvoice(errors=[f"Error: {str(e)}"]).to_dict()
    
    def process_pdf(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """Extract invoice data from PDF using multiple strategies."""
        invoice = ExtractedInvoice(source="pdf")
        strategies = []
        if PDFPLUMBER_AVAILABLE: strategies.append(self._extract_pdfplumber)
        if PYPDF2_AVAILABLE: strategies.append(self._extract_pypdf2)
        if not strategies:
            invoice.errors.append("No PDF libs available")
            return invoice.to_dict()
        best_text = ""
        for strat in strategies:
            try:
                text = strat(file_path)
                if len(text) > len(best_text): best_text = text
            except Exception as e:
                invoice.errors.append(f"{strat.__name__}: {str(e)}")
        if best_text:
            invoice.raw_text = best_text
            self._parse_data(invoice)
        else:
            invoice.errors.append("No text extracted")
        return invoice.to_dict()
    
    def _extract_pdfplumber(self, path: Path) -> str:
        parts = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                txt = page.extract_text()
                if txt: parts.append(txt)
        return "\n".join(parts)
    
    def _extract_pypdf2(self, path: Path) -> str:
        parts = []
        with open(path, 'rb') as f:
            for page in PyPDF2.PdfReader(f).pages:
                txt = page.extract_text()
                if txt: parts.append(txt)
        return "\n".join(parts)
    
    def process_image(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """Extract invoice data from image using OCR."""
        invoice = ExtractedInvoice(source="image")
        if not TESSERACT_AVAILABLE:
            invoice.errors.append("Tesseract not available")
            return invoice.to_dict()
        try:
            img = Image.open(file_path)
            if img.mode not in ['L', 'RGB']: img = img.convert('RGB')
            configs = ['--psm 6', '--psm 3', '--psm 4']
            best_text = ""
            for cfg in configs:
                try:
                    txt = pytesseract.image_to_string(img, config=cfg)
                    if len(txt) > len(best_text): best_text = txt
                except Exception as e:
                    invoice.errors.append(f"OCR {cfg}: {str(e)}")
            if best_text:
                invoice.raw_text = best_text
                self._parse_data(invoice)
            else:
                invoice.errors.append("No OCR text")
        except Exception as e:
            invoice.errors.append(f"Image error: {str(e)}")
        return invoice.to_dict()
    
    def process_docx(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """Extract invoice data from DOCX file."""
        invoice = ExtractedInvoice(source="docx")
        if not DOCX_AVAILABLE:
            invoice.errors.append("python-docx not available")
            return invoice.to_dict()
        try:
            doc = Document(file_path)
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells: parts.append(" | ".join(cells))
            invoice.raw_text = "\n".join(parts)
            self._parse_data(invoice)
        except Exception as e:
            invoice.errors.append(f"DOCX error: {str(e)}")
        return invoice.to_dict()
    
    def _parse_data(self, inv: ExtractedInvoice) -> None:
        """Parse structured data from raw text."""
        text = inv.raw_text
        for pat in self.PATTERNS["invoice_number"]:
            m = re.search(pat, text, re.I)
            if m: inv.invoice_number = m.group(1).strip(); break
        dates = re.findall(self.PATTERNS["date"][0], text)
        if len(dates) >= 1: inv.invoice_date = dates[0]
        if len(dates) >= 2: inv.due_date = dates[1]
        amounts = []
        for pat in self.PATTERNS["amount"]:
            for m in re.finditer(pat, text, re.I):
                try: amounts.append(float(m.group(1).replace(',', '')))
                except: pass
        if amounts:
            amounts = sorted(set(amounts), reverse=True)
            inv.total_amount = amounts[0]
            if len(amounts) > 1: inv.subtotal = amounts[1]
        for pat in self.PATTERNS["currency"]:
            m = re.search(pat, text)
            if m: inv.currency = m.group(0); break
        lines = [l.strip() for l in text.split('\n') if l.strip()]
        if lines:
            for line in lines[:10]:
                if re.search(r'(inc|llc|ltd|corp|co)\.?', line, re.I):
                    inv.vendor_name = line.strip(); break
            if not inv.vendor_name: inv.vendor_name = lines[0][:100]
        inv.line_items = self._extract_items(text)
        filled = sum([inv.invoice_number is not None, inv.invoice_date is not None,
                      inv.total_amount is not None, inv.vendor_name is not None])
        inv.confidence = filled / 4.0
    
    def _extract_items(self, text: str) -> List[Dict[str, Any]]:
        """Extract line items from invoice text."""
        items = []
        pat = r'(.*?)(\d+)\s*(?:x|@)?\s*[$€£¥]?\s*([\d,]+\.?\d{0,2})'
        for line in text.split('\n'):
            line = line.strip()
            if len(line) < 10: continue
            m = re.search(pat, line, re.I)
            if m:
                try:
                    items.append({
                        "description": m.group(1).strip()[:200],
                        "quantity": int(m.group(2)),
                        "unit_price": float(m.group(3).replace(',', '')),
                        "total": int(m.group(2)) * float(m.group(3).replace(',', ''))
                    })
                except: pass
        return items[:20]


def process_document(file_path: Union[str, Path], **kwargs) -> Dict[str, Any]:
    """Process a document and return extracted invoice data."""
    return DocumentProcessor(**kwargs).process(file_path)


if __name__ == "__main__":
    import sys, json
    if len(sys.argv) > 1:
        print(json.dumps(process_document(sys.argv[1]), indent=2, default=str))
    else:
        print("Usage: python document_processor.py <file_path>")
